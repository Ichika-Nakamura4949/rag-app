"""ドキュメント取り込みパイプライン."""

import base64
import logging
import shutil
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path

import chromadb
from langchain_chroma import Chroma
from langchain_core.documents import Document
from langchain_openai import OpenAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from openai import OpenAI

from app.core.config import Settings

logger = logging.getLogger(__name__)

CAPTION_PROMPT = """\
You are an assistant that describes images extracted from technical documents.
These descriptions will be stored in a vector database for semantic search.

Describe this image thoroughly:
- What type of visual is it? (diagram, chart, table, screenshot, flowchart, etc.)
- Transcribe ALL text visible in the image verbatim
- Describe the structure, relationships, and data shown
- Note key data points, numbers, and statistics
- Summarize the main concept being illustrated

Use specific terminology. Be detailed enough that someone could understand
the image content from your description alone."""


class IngestionPipeline:
    """PDF/DOCX解析 → チャンク分割 → Embedding → ChromaDB保存."""

    def __init__(self, settings: Settings) -> None:
        self._settings = settings
        self._embeddings = OpenAIEmbeddings(
            model=settings.embedding_model,
            openai_api_key=settings.openai_api_key,
        )
        self._splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.chunk_size,
            chunk_overlap=settings.chunk_overlap,
        )
        self._client = chromadb.PersistentClient(
            path=str(settings.chroma_persist_path),
        )
        self._openai = OpenAI(api_key=settings.openai_api_key)

    def _parse_pdf(self, file_path: Path) -> str:
        """PDFファイルからテキストを抽出."""
        import pymupdf

        doc = pymupdf.open(str(file_path))
        pages: list[str] = []
        for page in doc:
            text = page.get_text()
            if text.strip():
                pages.append(text)
        doc.close()
        return "\n\n".join(pages)

    def _parse_docx(self, file_path: Path) -> str:
        """DOCXファイルからテキストを抽出."""
        from docx import Document as DocxDocument

        doc = DocxDocument(str(file_path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)

    def _parse_file(self, file_path: Path) -> str:
        """ファイル形式に応じてテキストを抽出."""
        suffix = file_path.suffix.lower()
        if suffix == ".pdf":
            return self._parse_pdf(file_path)
        if suffix == ".docx":
            return self._parse_docx(file_path)
        raise ValueError(f"サポートされていないファイル形式: {suffix}")

    def _extract_images_from_pdf(
        self, file_path: Path, document_id: str
    ) -> list[tuple[Path, int]]:
        """PDFから画像を抽出してディスクに保存.

        Args:
            file_path: PDFファイルのパス
            document_id: ドキュメントID

        Returns:
            (画像パス, ページ番号) のリスト
        """
        import pymupdf

        doc = pymupdf.open(str(file_path))
        image_dir = self._settings.image_path / document_id
        image_dir.mkdir(parents=True, exist_ok=True)

        extracted: list[tuple[Path, int]] = []
        for page_num, page in enumerate(doc):
            image_list = page.get_images(full=True)
            img_index = 0
            for img_info in image_list:
                xref = img_info[0]
                base_image = doc.extract_image(xref)
                if not base_image:
                    continue

                width = base_image["width"]
                height = base_image["height"]

                # 小さいアイコン・装飾を除外
                if (
                    width < self._settings.min_image_width
                    or height < self._settings.min_image_height
                ):
                    continue

                ext = base_image["ext"]
                image_bytes = base_image["image"]
                image_path = image_dir / f"page{page_num}_img{img_index}.{ext}"
                image_path.write_bytes(image_bytes)

                extracted.append((image_path, page_num))
                img_index += 1
                logger.info(
                    "画像抽出: %s (page=%d, %dx%d)",
                    image_path.name,
                    page_num,
                    width,
                    height,
                )

        doc.close()
        logger.info(
            "画像抽出完了: document_id=%s, 画像数=%d", document_id, len(extracted)
        )
        return extracted

    def _caption_image(self, image_path: Path) -> str:
        """GPT-4o-miniで画像のキャプション（テキスト説明文）を生成.

        Args:
            image_path: 画像ファイルのパス

        Returns:
            画像の詳細な説明文
        """
        image_data = base64.b64encode(image_path.read_bytes()).decode("utf-8")
        ext = image_path.suffix.lstrip(".")
        media_type = f"image/{'jpeg' if ext in ('jpg', 'jpeg') else ext}"

        response = self._openai.chat.completions.create(
            model=self._settings.caption_model,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": CAPTION_PROMPT},
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:{media_type};base64,{image_data}",
                            },
                        },
                    ],
                }
            ],
            max_tokens=1000,
        )
        caption = response.choices[0].message.content or ""
        logger.info("キャプション生成: %s (%d文字)", image_path.name, len(caption))
        return caption

    def ingest(
        self, file_path: Path, document_id: str, filename: str
    ) -> tuple[int, int]:
        """ドキュメントを解析してChromaDBに保存.

        Args:
            file_path: ドキュメントファイルのパス
            document_id: ドキュメントの一意ID
            filename: 元のファイル名

        Returns:
            (テキストチャンク数, 画像数) のタプル
        """
        # テキスト抽出・チャンク分割
        text = self._parse_file(file_path)
        documents: list[Document] = []

        if text.strip():
            chunks = self._splitter.split_text(text)
            documents.extend(
                Document(
                    page_content=chunk,
                    metadata={
                        "document_id": document_id,
                        "filename": filename,
                        "chunk_index": i,
                        "type": "text",
                    },
                )
                for i, chunk in enumerate(chunks)
            )

        text_chunk_count = len(documents)

        # PDF画像抽出・キャプション並列生成
        image_count = 0
        if file_path.suffix.lower() == ".pdf":
            extracted_images = self._extract_images_from_pdf(file_path, document_id)
            if extracted_images:
                logger.info(
                    "キャプション並列生成開始: %d枚", len(extracted_images)
                )
                with ThreadPoolExecutor(max_workers=5) as executor:
                    future_to_image = {
                        executor.submit(self._caption_image, img_path): (
                            img_path,
                            page_num,
                        )
                        for img_path, page_num in extracted_images
                    }
                    for future in as_completed(future_to_image):
                        img_path, page_num = future_to_image[future]
                        try:
                            caption = future.result()
                            if caption.strip():
                                relative_path = (
                                    f"{document_id}/{img_path.name}"
                                )
                                documents.append(
                                    Document(
                                        page_content=caption,
                                        metadata={
                                            "document_id": document_id,
                                            "filename": filename,
                                            "chunk_index": text_chunk_count
                                            + image_count,
                                            "type": "image_caption",
                                            "image_path": relative_path,
                                            "page_number": page_num,
                                        },
                                    )
                                )
                                image_count += 1
                        except Exception:
                            logger.exception(
                                "キャプション生成に失敗: %s", img_path.name
                            )

        # ChromaDBに保存
        if documents:
            vectorstore = Chroma(
                client=self._client,
                collection_name="documents",
                embedding_function=self._embeddings,
            )
            vectorstore.add_documents(documents)

        logger.info(
            "Ingestion完了: document_id=%s, テキストチャンク=%d, 画像=%d",
            document_id,
            text_chunk_count,
            image_count,
        )
        return text_chunk_count, image_count

    def delete_by_document_id(self, document_id: str) -> None:
        """指定ドキュメントIDのチャンクをChromaDBから削除."""
        collection = self._client.get_or_create_collection("documents")
        collection.delete(where={"document_id": document_id})

        # 画像ディレクトリも削除
        image_dir = self._settings.image_path / document_id
        if image_dir.exists():
            shutil.rmtree(image_dir)
            logger.info("画像ディレクトリ削除: %s", image_dir)
